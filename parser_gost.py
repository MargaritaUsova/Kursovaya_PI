import re
from docx import Document
from bs4 import BeautifulSoup
import requests

#document = Document('TestKURS.docx')
#document.save('TestKURS.docx')
def parser(doc):

    document = Document(doc)

    def OnlyRus():
        return re.compile("aDocIntro+\d[0-9]")

    gosts = []
    art = []
    dict_gost = {}
    res = []
    #for count in range(1,2437):
    #тест для первых 200
    for count in range(1, 200): 
        url = f'https://www.gostinfo.ru/catalog/gostlist?page={count}'
        page = requests.get(url)
        soup = BeautifulSoup(page.text, "html.parser")
        gosts = gosts + (soup.findAll('span', class_='textBlue'))
        art = art + (soup.find_all('a', id = OnlyRus()))

    #делаю словарь 
    for i in range(len(gosts)):
        dict_gost[gosts[i].get_text()] = art[i].get_text()

    def in_dictionary(dict):  
        k = []
        v = []
        for key in dict:
            k.append(key)
            v.append(dict[key])
        return k, v

    k, v = in_dictionary(dict_gost)

    for para in document.paragraphs:
        for i in range(len(dict_gost)):
            if k[i] in para.text and v[i] in para.text:
                res.append(k[i]+" "+v[i])

    return res
    #Все Госты, которые встр на первых 200 стр сайта и в документе. 