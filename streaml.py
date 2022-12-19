import streamlit as st
import requests
import os.path
import pathlib
import re
from docx import Document
from bs4 import BeautifulSoup
import requests

#ocument = Document('TestKURS.docx')
#document.save('TestKURS.docx')


uploaded_file = st.file_uploader("Upload a file")

if uploaded_file:
   st.write("Filename: ", uploaded_file.name)


document = Document(uploaded_file)

def OnlyRus():
  return re.compile("aDocIntro+\d[0-9]")

gosts = []
art = []
dict_gost = {}
#for count in range(1,2437):
#тест для первых 200
for count in range(200): 
    url = f'https://www.gostinfo.ru/catalog/gostlist?page={count}'
    page = requests.get(url)
    soup = BeautifulSoup(page.text, "html.parser")
    gosts = gosts + (soup.findAll('span', class_='textBlue'))
    art = art + (soup.find_all('a', id = OnlyRus()))
    
    #art = soup.findAll('a')
    #gosts.text()
    #for gost in gosts:
      #print(gost.get_text())

#делаю словарь 
for i in range(len(gosts)):
  dict_gost[gosts[i].get_text()] = art[i].get_text()
#print(dict_gost)


def in_dictionary(dict):  
  k = []
  v = []
  for key in dict:
    k.append(key)
    v.append(dict[key])
  return k, v

k, v = in_dictionary(dict_gost)

for para in document.paragraphs:
  for i in range(len(dict_gost)):
    if k[i] in para.text and v[i] in para.text:
      st.write(k[i]+v[i])


